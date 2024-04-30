import os
import tkinter as tk
from tkinter import ttk, filedialog
from docx import Document
from docx2pdf import convert
import tempfile
import re

# Run the following command in your terminal while in the project directory:
# pip install python-docx
# pip install docx2pdf

# After installing dependencies run the app by typing in your terminal:
# python main.py

class App():
    def __init__(self):
        self.labels = []
        self.entry_boxes = []

        self.root = tk.Tk()
        self.root.title('Text Replacement Automation Tool')
        self.mainframe = tk.Frame(self.root)
        self.mainframe.pack(fill='both', expand=True)
        
        # Source File Section
        self.source_frame = ttk.LabelFrame(self.mainframe, text="Source File")
        self.source_frame.grid(row=0, column=0, sticky='NWES', padx=20, pady=10)
        select_file_btn = ttk.Button(self.source_frame, text='Select File', command=self.select_file)
        select_file_btn.grid(row=0, column=0, padx=20, pady=10)
        self.source_file = ttk.Label(self.source_frame, text='')
        self.source_file.grid(row=0, column=2, padx=20, pady=10)
        
        # Replace Text Section
        self.re_frame = ttk.LabelFrame(self.mainframe, text="Replace Text")
        self.re_frame.grid(row=1, column=0, sticky='NWES', padx=20, pady=10)
        
        for widget in self.re_frame.winfo_children():
            widget.grid_configure(padx=10,pady=5)

        # Export Section
        self.finalize = ttk.Labelframe(self.mainframe, text="Export")
        self.finalize.grid(row=2, column=0, sticky="NWES", padx=20, pady=10)
        select_file_type_label = ttk.Label(self.finalize, text='File Extension Type:')
        select_file_type_label.grid(row=1, column=0)
        self.select_file_type = ttk.Combobox(self.finalize, values=['.docx','.pdf'])
        self.select_file_type.grid(row=1, column=1)
        export_btn = ttk.Button(self.finalize, text='Save as', command=self.export)
        export_btn.grid(row=2, column=1)

        for widget in self.finalize.winfo_children():
            widget.grid_configure(padx=20,pady=5)
        
        # Instructions Section
        self.instruction_frame = ttk.LabelFrame(self.mainframe, text="Instructions")
        self.instruction_frame.grid(row=3, column=0, sticky="NWES", padx=20, pady=10)
        step_1_label = ttk.Label(self.instruction_frame, text='  Step 1: Select a .docx file that has text within square brackets to be replaced ex. [text]')
        step_1_label.grid(row=0, column=0, sticky='w', padx=20, pady=5)
        step_2_label = ttk.Label(self.instruction_frame, text='  Step 2: Fill out the generated fields to replace the corresponding text')
        step_2_label.grid(row=1, column=0, sticky='w', padx=20, pady=5)
        step_3_label = ttk.Label(self.instruction_frame, text='  Step 3: When finished, select a file extension type, click save as, and name your file')
        step_3_label.grid(row=2, column=0, sticky='w', padx=20, pady=5)
        important_note = ttk.Label(self.instruction_frame, text='Please Note:')
        important_note.grid(row=3, column=0, sticky='w', padx=20, pady=5)
        note_1 = ttk.Label(self.instruction_frame, text='  - Must import using .docx file')
        note_1.grid(row=4, column=0, sticky='w', padx=20, pady=5)
        note_2 = ttk.Label(self.instruction_frame, text='  - If exporting to .pdf a popup will appear, click ok on the default option')
        note_2.grid(row=5, column=0, sticky='w', padx=20, pady=5)
         
        self.root.mainloop()
        return
    
    def select_file(self):
        self.file_path = filedialog.askopenfilename()
        file_name = os.path.basename(self.file_path)
        self.source_file.config(text=file_name)
        print('Selected file:', self.file_path)
        self.replace_labels()

    def replace_labels(self):
        self.labels = []
        doc = Document(self.file_path)
        for paragraph in doc.paragraphs:
            # regex to find text within square brackets (change [] if you want to change delimiter)
            matches = re.findall(r'(\[.*?\])', paragraph.text)
            for match in matches:
                if match not in self.labels:
                    self.labels.append(match)
        if self.labels:
            print("Labels found:")
            for label in self.labels:
                print(label)
        else:
            print("No labels found in the document.")
        self.render_labels()


    def render_labels(self):
        for widget in self.re_frame.winfo_children():
            widget.destroy()

        # dynamically render labels based on self.labels
        num_row = 0
        for i, label_text in enumerate(self.labels):
            if i != 0 and (i % 3) == 0:
                num_row += 2
            label = ttk.Label(self.re_frame, text=label_text)
            label.grid(row=num_row, column=(i % 3), sticky='NWES')
            entry = ttk.Entry(self.re_frame)
            entry.grid(row=num_row + 1, column=(i % 3), sticky='NWES')
            self.entry_boxes.append(entry)

        for widget in self.re_frame.winfo_children():
            widget.grid_configure(padx=10,pady=5)

    def export(self):
        if not hasattr(self, 'file_path') or not self.file_path:
            print('No file selected!')
            return
        if not self.select_file_type.get():
            print('No file extension selected!')
            return

        print(self.select_file_type.get())
        if self.select_file_type.get() == '.docx':
            self.export_docx()
        elif self.select_file_type.get() == '.pdf':
            self.export_pdf()
        else:
            print(f"Unsupported file extension for export: {self.select_file_type.get()}")
            return

    def export_docx(self):
        doc = Document(self.file_path)
        for label, entry in zip(self.labels, self.entry_boxes):
            keyword = label
            new_value = entry.get()
            for paragraph in doc.paragraphs:
                if keyword in paragraph.text:
                    paragraph.text = paragraph.text.replace(keyword, new_value)
        new_file = filedialog.asksaveasfilename(defaultextension='.docx')
        if not new_file:
            print('Export canceled!')
            return
        doc.save(new_file)
        print('File exported successfully!')

    def export_pdf(self):
        if not hasattr(self, 'file_path') or not self.file_path:
            print("No file selected!")
            return
        try:
            doc = Document(self.file_path)
        except Exception as e:
            print(f"Error reading {self.file_path}: {e}")
            return
        # Replace words in document
        for paragraph in doc.paragraphs:
            for label, entry in zip(self.labels, self.entry_boxes):
                keyword = label
                new_value = entry.get()
                if keyword in paragraph.text:
                    paragraph.text = paragraph.text.replace(keyword, new_value)
        new_file = filedialog.asksaveasfilename(defaultextension=".pdf")
        if not new_file:
            print("Export canceled.")
            return
        # Save modified document to temp file
        temp_docx_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        try:
            doc.save(temp_docx_file.name)
        except Exception as e:
            print(f"Error saving temporary DOCX file: {e}")
            return
        # Convert temp docx file to pdf using docx2pdf
        try:
            convert(temp_docx_file.name, new_file)
            print(f"File exported successfully: {new_file}")
        except Exception as e:
            print(f"Error converting to PDF: {e}")
        finally:
            temp_docx_file.close()
            os.unlink(temp_docx_file.name)

if __name__ == '__main__':
    App()